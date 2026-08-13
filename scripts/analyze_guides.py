from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
GUIDE_DIR = ROOT / "guide"
OUTPUT = ROOT / "tmp" / "guide-analysis.json"


def volume_number(path: Path) -> int:
    match = re.search(r"Volume_(\d+)", path.name)
    return int(match.group(1)) if match else 999


def paragraph_hyperlinks(paragraph) -> list[dict[str, str]]:
    links: list[dict[str, str]] = []
    rels = paragraph.part.rels
    for hyperlink in paragraph._p.findall(qn("w:hyperlink")):
        rel_id = hyperlink.get(qn("r:id"))
        text = "".join(node.text or "" for node in hyperlink.iter(qn("w:t")))
        if rel_id and rel_id in rels:
            links.append({"text": text.strip(), "url": rels[rel_id].target_ref})
    return links


def analyze(path: Path) -> dict:
    document = Document(path)
    paragraphs = []
    links = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        paragraph_links = paragraph_hyperlinks(paragraph)
        links.extend(paragraph_links)
        if text:
            paragraphs.append({"style": paragraph.style.name, "text": text})

    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                rows.append(values)
        if rows:
            tables.append(rows)

    return {
        "volume": volume_number(path),
        "filename": path.name,
        "title": document.core_properties.title,
        "paragraphs": paragraphs,
        "tables": tables,
        "links": links,
        "inline_shapes": len(document.inline_shapes),
    }


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    documents = sorted(GUIDE_DIR.glob("*.docx"), key=volume_number)
    OUTPUT.write_text(
        json.dumps([analyze(path) for path in documents], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(OUTPUT)


if __name__ == "__main__":
    main()
